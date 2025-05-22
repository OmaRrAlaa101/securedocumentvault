import os
import re
import bcrypt
from datetime import datetime, timedelta
from flask import Flask, render_template, request, redirect, url_for, flash, session, abort, make_response, send_file
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager, login_user, logout_user, login_required, UserMixin, current_user
from authlib.integrations.flask_client import OAuth
from sqlalchemy.exc import IntegrityError
from functools import wraps
import pyotp
from cryptography.fernet import Fernet
from werkzeug.utils import secure_filename
import hashlib
import hmac
from pathlib import Path
try:
    import magic
except ImportError:
    magic = None
    print("python-magic is not installed. Please install it with 'pip install python-magic python-magic-bin'.")
import jwt
from dotenv import load_dotenv
import qrcode
import io
import base64
import random
import string
import zlib
from cryptography.hazmat.primitives import hashes, serialization
from cryptography.hazmat.primitives.asymmetric import padding, rsa
from cryptography.hazmat.primitives.asymmetric import utils as asym_utils
from cryptography.hazmat.backends import default_backend
from flask_mail import Mail, Message
import secrets

# Load environment variables from .env file
load_dotenv()
print('GITHUB_CLIENT_ID:', os.environ.get('GITHUB_CLIENT_ID'))
print('GITHUB_CLIENT_SECRET:', os.environ.get('GITHUB_CLIENT_SECRET'))

def generate_secret_key(length=32):
    return ''.join(random.choices(string.ascii_letters + string.digits, k=length))

# --- Flask App Setup ---
app = Flask(__name__)
if app.debug:
    # For development, generate a new key every time (forces logout on restart)
    app.secret_key = generate_secret_key()
else:
    # In production, use a fixed key from environment
    app.secret_key = os.environ.get("SECRET_KEY", "supersecretkey")
app.config['SQLALCHEMY_DATABASE_URI'] = f'sqlite:///{os.path.join(os.path.dirname(os.path.abspath(__file__)), "auth_system.db")}'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['REMEMBER_COOKIE_DURATION'] = timedelta(days=7)
app.config['UPLOAD_FOLDER'] = os.path.join(os.path.dirname(os.path.abspath(__file__)), "secure_uploads")
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# Enforce HTTPS in production
if not app.debug:
    app.config['SESSION_COOKIE_SECURE'] = True
    app.config['REMEMBER_COOKIE_SECURE'] = True

    @app.before_request
    def before_request_https_redirect():
        if not request.is_secure and request.headers.get('X-Forwarded-Proto', 'http') != 'https':
            url = request.url.replace('http://', 'https://', 1)
            return redirect(url, code=301)

# Ensure upload directory exists
Path(app.config['UPLOAD_FOLDER']).mkdir(parents=True, exist_ok=True)

# Initialize encryption key
if not os.path.exists('encryption.key'):
    key = Fernet.generate_key()
    with open('encryption.key', 'wb') as key_file:
        key_file.write(key)

with open('encryption.key', 'rb') as key_file:
    ENCRYPTION_KEY = key_file.read()
fernet = Fernet(ENCRYPTION_KEY)

# --- Database Setup ---
db = SQLAlchemy(app)
login_manager = LoginManager(app)
login_manager.login_view = 'login'

# --- OAuth Setup ---
oauth = OAuth(app)
github = oauth.register(
    name='github',
    client_id=os.environ.get('GITHUB_CLIENT_ID'),
    client_secret=os.environ.get('GITHUB_CLIENT_SECRET'),
    access_token_url='https://github.com/login/oauth/access_token',
    access_token_params=None,
    authorize_url='https://github.com/login/oauth/authorize',
    authorize_params=None,
    api_base_url='https://api.github.com/',
    client_kwargs={'scope': 'user:email'},
)

google = oauth.register(
    name='google',
    client_id=os.environ.get('GOOGLE_CLIENT_ID'),
    client_secret=os.environ.get('GOOGLE_CLIENT_SECRET'),
    access_token_url='https://oauth2.googleapis.com/token',
    access_token_params=None,
    authorize_url='https://accounts.google.com/o/oauth2/auth',
    authorize_params=None,
    api_base_url='https://www.googleapis.com/oauth2/v1/',
    client_kwargs={'scope': 'email profile'},
)

# --- Models ---
class Role(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(80), unique=True)
    users = db.relationship('User', backref='role', lazy=True)

class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), unique=True, nullable=True)
    email = db.Column(db.String(120), unique=True, nullable=True)
    password_hash = db.Column(db.String(128), nullable=True)
    github_id = db.Column(db.String(64), unique=True, nullable=True)
    auth_method = db.Column(db.String(20), nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    role_id = db.Column(db.Integer, db.ForeignKey('role.id'))
    two_factor_secret = db.Column(db.String(32))
    two_factor_enabled = db.Column(db.Boolean, default=False)
    documents = db.relationship('Document', backref='owner', lazy=True)
    login_logs = db.relationship('LoginLog', backref='user', lazy=True)
    phone = db.Column(db.String(32), nullable=True)
    department = db.Column(db.String(64), nullable=True)
    job_title = db.Column(db.String(64), nullable=True)
    profile_pic_url = db.Column(db.String(255), nullable=True)

    def get_2fa_uri(self):
        if self.two_factor_secret:
            return pyotp.totp.TOTP(self.two_factor_secret).provisioning_uri(
                name=self.email,
                issuer_name="SecureDocs"
            )
        return None

    def verify_2fa(self, token):
        if not self.two_factor_secret:
            return False
        totp = pyotp.TOTP(self.two_factor_secret)
        return totp.verify(token)

    def is_admin(self):
        return self.role.name == 'admin' if self.role else False

class Document(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    filename = db.Column(db.String(255), nullable=False)
    original_filename = db.Column(db.String(255), nullable=False)
    file_type = db.Column(db.String(64), nullable=False)
    file_size = db.Column(db.Integer, nullable=False)
    upload_date = db.Column(db.DateTime, default=datetime.utcnow)
    last_modified = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    hash_sha256 = db.Column(db.String(64), nullable=False)
    signature = db.Column(db.Text, nullable=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    is_encrypted = db.Column(db.Boolean, default=True)
    hmac = db.Column(db.String(64), nullable=False)
    crc32 = db.Column(db.String(16), nullable=True)
    openssl_signature = db.Column(db.Text, nullable=True)
    description = db.Column(db.String(255), nullable=True)

class LoginLog(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=True)
    timestamp = db.Column(db.DateTime, default=datetime.utcnow)
    ip_address = db.Column(db.String(45))
    user_agent = db.Column(db.String(256))
    method = db.Column(db.String(20))
    success = db.Column(db.Boolean)
    details = db.Column(db.String(256))

# --- Helper Functions ---
def calculate_file_hash(file_path):
    sha256_hash = hashlib.sha256()
    with open(file_path, "rb") as f:
        for byte_block in iter(lambda: f.read(4096), b""):
            sha256_hash.update(byte_block)
    return sha256_hash.hexdigest()

def calculate_hmac(file_path):
    hmac_obj = hmac.new(ENCRYPTION_KEY, digestmod=hashlib.sha256)
    with open(file_path, "rb") as f:
        for byte_block in iter(lambda: f.read(4096), b""):
            hmac_obj.update(byte_block)
    return hmac_obj.hexdigest()

def encrypt_file(file_path):
    with open(file_path, 'rb') as file:
        file_data = file.read()
    encrypted_data = fernet.encrypt(file_data)
    with open(file_path, 'wb') as file:
        file.write(encrypted_data)

def decrypt_file(file_path):
    with open(file_path, 'rb') as file:
        encrypted_data = file.read()
    decrypted_data = fernet.decrypt(encrypted_data)
    return decrypted_data

def allowed_file(filename):
    ALLOWED_EXTENSIONS = {'txt', 'pdf', 'doc', 'docx'}
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# --- Decorators ---
def admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not current_user.is_authenticated or not current_user.is_admin():
            abort(403)
        return f(*args, **kwargs)
    return decorated_function

def require_2fa(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if current_user.two_factor_enabled and not session.get('2fa_verified'):
            return redirect(url_for('verify_2fa'))
        return f(*args, **kwargs)
    return decorated_function

# --- User Loader ---
@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

# --- Password Policy ---
def valid_password(password):
    # At least 8 characters, one uppercase, one lowercase, one digit, one special char
    if (len(password) < 8 or
        not re.search(r'[A-Z]', password) or
        not re.search(r'[a-z]', password) or
        not re.search(r'[0-9]', password) or
        not re.search(r'[^A-Za-z0-9]', password)):
        return False
    return True

# --- Routes ---
@app.route('/')
@login_required
def home():
    # Show all login logs for the current user
    recent_logs = LoginLog.query.filter_by(user_id=current_user.id)\
        .order_by(LoginLog.timestamp.desc())\
        .all()
    return render_template('home.html', user=current_user, recent_logs=recent_logs)

def save_user_to_file(user):
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"database_export_{timestamp}.txt"
    
    with open(filename, 'w') as f:
        f.write("=== New User Registration ===\n\n")
        f.write(f"User ID: {user.id}\n")
        f.write(f"Username: {user.username}\n")
        f.write(f"Email: {user.email}\n")
        f.write(f"Auth Method: {user.auth_method}\n")
        f.write(f"Created At: {user.created_at}\n")
        f.write("\n" + "="*50 + "\n")

@app.route('/signup', methods=['GET', 'POST'])
def signup():
    if request.method == 'POST':
        username = request.form['username'].strip()
        email = request.form['email'].strip().lower()
        password = request.form['password']
        is_admin_checked = 'is_admin' in request.form

        if not username or not email or not password:
            flash('All fields are required.', 'danger')
            return redirect(url_for('signup'))
        if not valid_password(password):
            flash('Password too weak. Must be at least 8 characters, include numbers and symbols.', 'danger')
            return redirect(url_for('signup'))
        if User.query.filter_by(email=email).first():
            flash('Email already used, try another.', 'danger')
            return redirect(url_for('signup'))
        hashed = bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt())

        admin_role = Role.query.filter_by(name='admin').first()
        user_role = Role.query.filter_by(name='user').first()
        user = User(
            username=username,
            email=email,
            password_hash=hashed.decode('utf-8'),
            auth_method='manual',
            role_id=admin_role.id if is_admin_checked else user_role.id
        )

        admin_exists = User.query.join(Role).filter(Role.name == 'admin').first() is not None
        if is_admin_checked and admin_exists:
            flash('An admin already exists. You cannot create another admin from signup.', 'danger')
            return redirect(url_for('signup'))

        try:
            db.session.add(user)
            db.session.commit()
            save_user_to_file(user)
            flash('Signup successful! Please set up 2FA.', 'success')
            login_user(user)  # Log in directly after signup
            return redirect(url_for('setup_2fa'))
        except IntegrityError:
            db.session.rollback()
            flash('Username or email already exists.', 'danger')
    return render_template('signup.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        method = 'manual'
        email = request.form['email'].strip().lower()
        password = request.form['password']
        remember = 'remember' in request.form
        user = User.query.filter_by(email=email, auth_method='manual').first()
        success = False
        details = ''
        if user and user.password_hash and bcrypt.checkpw(password.encode('utf-8'), user.password_hash.encode('utf-8')):
            login_user(user, remember=remember)
            success = True
            flash('Logged in successfully.', 'success')
            resp = make_response(redirect(url_for('post_login_2fa')))
            resp.headers['Cache-Control'] = 'no-store'
            log_login(user.id, method, success, request, details)
            return resp
        else:
            details = 'Invalid credentials'
            flash('Invalid email or password, please try again.', 'danger')
            log_login(user.id if user else None, method, success, request, details)
    return render_template('login.html')

@app.route('/post-login-2fa')
@login_required
def post_login_2fa():
    if current_user.two_factor_enabled and not session.get('2fa_verified'):
        return redirect(url_for('verify_2fa'))
    elif not current_user.two_factor_enabled:
        return redirect(url_for('setup_2fa'))
    return redirect(url_for('home'))

@app.route('/login/github')
def github_login():
    # Check for GitHub OAuth credentials
    if not os.environ.get('GITHUB_CLIENT_ID') or not os.environ.get('GITHUB_CLIENT_SECRET'):
        flash('GitHub OAuth credentials are not set. Please set GITHUB_CLIENT_ID and GITHUB_CLIENT_SECRET in your .env file.', 'danger')
        return redirect(url_for('login'))
    session['next_url'] = request.args.get('next') or url_for('home')
    redirect_uri = url_for('github_authorized', _external=True)
    return github.authorize_redirect(redirect_uri)

@app.route('/logout')
@login_required
def logout():
    # Clear all session data
    session.clear()
    
    # Clear OAuth tokens if they exist
    if 'github_token' in session:
        del session['github_token']
    if 'google_token' in session:
        del session['google_token']
    
    # Clear 2FA verification
    if '2fa_verified' in session:
        del session['2fa_verified']
    
    # Logout the user
    logout_user()
    
    # Clear remember me cookie
    resp = make_response(redirect(url_for('login')))
    resp.delete_cookie('remember_token')
    
    # Set cache control headers
    resp.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, post-check=0, pre-check=0, max-age=0'
    resp.headers['Pragma'] = 'no-cache'
    resp.headers['Expires'] = '0'
    
    flash('You have been logged out successfully.', 'info')
    return resp

@app.route('/login/github/authorized')
def github_authorized():
    try:
        token = github.authorize_access_token()
        if token is None or 'access_token' not in token:
            flash('Access denied: reason={} error={}'.format(
                request.args.get('error', 'unknown'),
                request.args.get('error_description', 'unknown')
            ), 'danger')
            return redirect(url_for('login'))

        session['github_token'] = token
        resp = github.get('user')
        if not resp or resp.status_code != 200:
            flash('Failed to get user info from GitHub.', 'danger')
            return redirect(url_for('login'))
        me = resp.json()
        username = me.get('login')
        email = me.get('email')

        emails_resp = github.get('user/emails')
        if emails_resp.status_code == 200:
            emails = emails_resp.json()
            for e in emails:
                if e.get('primary') and e.get('verified'):
                    email = e.get('email')
                    break
            if not email:
                for e in emails:
                    if e.get('verified'):
                        email = e.get('email')
                        break
            if not email and emails:
                email = emails[0].get('email')
        if not email:
            print('DEBUG: GitHub user info:', me)
            print('DEBUG: GitHub emails:', emails_resp.json() if emails_resp.status_code == 200 else emails_resp.text)
            flash('Could not get email from GitHub. Please make sure your GitHub account has a public or verified email.', 'danger')
            return redirect(url_for('login'))

        # --- FIX: If user with this email exists, just log them in ---
        user = User.query.filter_by(email=email).first()
        if not user:
            user_role = Role.query.filter_by(name='user').first()
            user = User(
                username=username,
                email=email,
                auth_method='github',
                role_id=user_role.id,
            )
            db.session.add(user)
            db.session.commit()
            save_user_to_file(user)
        login_user(user, remember=True)
        flash('Logged in successfully with GitHub.', 'success')
        next_url = session.pop('next_url', url_for('home'))
        resp = make_response(redirect(next_url))
        resp.headers['Cache-Control'] = 'no-store'
        return resp
    except Exception as e:
        import traceback
        print('GitHub login error:', e)
        traceback.print_exc()
        flash('An error occurred during GitHub authentication.', 'danger')
        return redirect(url_for('login'))

@app.route('/login/google')
def google_login():
    redirect_uri = url_for('google_authorized', _external=True)
    return google.authorize_redirect(redirect_uri)

@app.route('/login/google/authorized')
def google_authorized():
    token = google.authorize_access_token()
    resp = google.get('userinfo')
    user_info = resp.json()
    email = user_info['email']
    username = user_info.get('name', email.split('@')[0])
    google_id = user_info['id']

    user = User.query.filter_by(email=email, auth_method='google').first()
    if not user:
        user_role = Role.query.filter_by(name='user').first()
        user = User(
            username=username,
            email=email,
            auth_method='google',
            role_id=user_role.id,
        )
        db.session.add(user)
        db.session.commit()
        save_user_to_file(user)
    login_user(user, remember=True)
    flash('Logged in successfully with Google.', 'success')
    return redirect(url_for('home'))

def log_login(user_id, method, success, req, details):
    log = LoginLog(
        user_id=user_id,
        ip_address=req.remote_addr,
        user_agent=req.headers.get('User-Agent'),
        method=method,
        success=success,
        details=details
    )
    db.session.add(log)
    db.session.commit()

# --- Prevent Back Navigation After Logout ---
@app.after_request
def add_header(response):
    response.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, post-check=0, pre-check=0, max-age=0'
    response.headers['Pragma'] = 'no-cache'
    return response

# --- Error Handlers ---
@app.errorhandler(404)
def not_found(e):
    return render_template('404.html'), 404

@app.errorhandler(500)
def server_error(e):
    return render_template('500.html'), 500

# --- CLI to create DB ---
@app.cli.command('init-db')
def init_db():
    db.create_all()
    print("Database initialized.")

@app.route('/manage-users')
@login_required
@admin_required
def manage_users():
    users = User.query.all()
    return render_template('manage_users.html', users=users)

@app.route('/delete-user/<int:user_id>', methods=['POST'])
@login_required
@admin_required
def delete_user(user_id):
    if user_id == current_user.id:
        flash('You cannot delete your own account.', 'danger')
        return redirect(url_for('manage_users'))
    
    user = User.query.get_or_404(user_id)
    try:
        # Delete all login logs for this user
        LoginLog.query.filter_by(user_id=user_id).delete()
        # Delete the user
        db.session.delete(user)
        db.session.commit()
        flash('User deleted successfully.', 'success')
    except Exception as e:
        db.session.rollback()
        flash('Error deleting user.', 'danger')
    
    return redirect(url_for('manage_users'))

@app.route('/setup-2fa')
@login_required
def setup_2fa():
    if not current_user.two_factor_secret:
        current_user.two_factor_secret = pyotp.random_base32()
        db.session.commit()
    qr_uri = current_user.get_2fa_uri()
    img = qrcode.make(qr_uri)
    buf = io.BytesIO()
    img.save(buf, format='PNG')
    img_b64 = base64.b64encode(buf.getvalue()).decode('utf-8')
    return render_template('setup_2fa.html', qr_code=img_b64)

@app.route('/verify-2fa', methods=['GET', 'POST'])
@login_required
def verify_2fa():
    if request.method == 'POST':
        token = request.form.get('token')
        if current_user.verify_2fa(token):
            session['2fa_verified'] = True
            return redirect(url_for('home'))
        flash('Invalid 2FA token', 'danger')
    return render_template('verify_2fa.html')

@app.route('/enable-2fa', methods=['POST'])
@login_required
def enable_2fa():
    token = request.form.get('token')
    if current_user.verify_2fa(token):
        current_user.two_factor_enabled = True
        db.session.commit()
        flash('2FA has been enabled', 'success')
        return redirect(url_for('home'))
    flash('Invalid token', 'danger')
    return redirect(url_for('setup_2fa'))

def get_file_type(filename, file_path):
    # Try to use magic if available
    if magic is not None:
        try:
            return magic.from_file(file_path, mime=True)
        except Exception:
            pass
    # Fallback to extension
    ext = filename.rsplit('.', 1)[-1].lower()
    if ext == 'pdf':
        return 'application/pdf'
    elif ext == 'txt':
        return 'text/plain'
    elif ext == 'docx':
        return 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    elif ext == 'doc':
        return 'application/msword'
    else:
        return ext

@app.route('/upload-document', methods=['GET', 'POST'])
@login_required
@require_2fa
def upload_document():
    if request.method == 'POST':
        if 'file' not in request.files:
            flash('No file part', 'danger')
            return redirect(request.url)
        file = request.files['file']
        description = request.form.get('description', '').strip()
        if file.filename == '':
            flash('No selected file', 'danger')
            return redirect(request.url)
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            unique_filename = f"{timestamp}_{filename}"
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
            file.save(file_path)
            file_size = os.path.getsize(file_path)
            file_type = get_file_type(filename, file_path)
            file_hash = calculate_file_hash(file_path)
            crc32_value = calculate_crc32(file_path)
            openssl_signature = sign_file_with_private_key(file_path)
            with open(file_path, 'rb') as f:
                sha256_signature = hashlib.sha256(f.read()).hexdigest()
            encrypt_file(file_path)
            hmac_value = calculate_hmac(file_path)
            document = Document(
                filename=unique_filename,
                original_filename=filename,
                file_type=file_type,
                file_size=file_size,
                hash_sha256=file_hash,
                user_id=current_user.id,
                hmac=hmac_value,
                crc32=crc32_value,
                openssl_signature=openssl_signature,
                signature=sha256_signature,
                description=description
            )
            db.session.add(document)
            db.session.commit()
            flash('Document uploaded successfully', 'success')
            return redirect(url_for('list_documents'))
        flash('File type not allowed', 'danger')
        return redirect(request.url)
    return render_template('upload_document.html')

@app.route('/documents')
@login_required
@require_2fa
def list_documents():
    if current_user.is_admin():
        documents = Document.query.all()
    else:
        documents = Document.query.filter_by(user_id=current_user.id).all()
    return render_template('documents.html', documents=documents)

@app.route('/download-document/<int:doc_id>')
@login_required
@require_2fa
def download_document(doc_id):
    document = Document.query.get_or_404(doc_id)
    
    # Check permissions
    if not current_user.is_admin() and document.user_id != current_user.id:
        abort(403)
    
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], document.filename)
    
    # Verify HMAC
    current_hmac = calculate_hmac(file_path)
    if current_hmac != document.hmac:
        flash('Document integrity check failed (HMAC)', 'danger')
        return redirect(url_for('list_documents'))
    # Decrypt the file
    decrypted_data = decrypt_file(file_path)
    # Verify CRC32 on decrypted data
    crc32_value = format(zlib.crc32(decrypted_data) & 0xFFFFFFFF, '08x')
    if document.crc32 and crc32_value != document.crc32:
        flash('Document integrity check failed (CRC32)', 'danger')
        return redirect(url_for('list_documents'))
    # Verify OpenSSL digital signature
    if document.openssl_signature and not verify_file_signature(decrypted_data, document.openssl_signature):
        flash('Document signature verification failed (OpenSSL)', 'danger')
        return redirect(url_for('list_documents'))
    # Create response
    response = make_response(decrypted_data)
    response.headers['Content-Type'] = document.file_type
    response.headers['Content-Disposition'] = f'attachment; filename={document.original_filename}'
    return response

@app.route('/delete-document/<int:doc_id>', methods=['POST'])
@login_required
@require_2fa
def delete_document(doc_id):
    document = Document.query.get_or_404(doc_id)
    
    # Check permissions
    if not current_user.is_admin() and document.user_id != current_user.id:
        abort(403)
    
    # Delete the file
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], document.filename)
    if os.path.exists(file_path):
        os.remove(file_path)
    
    # Delete the record
    db.session.delete(document)
    db.session.commit()
    
    flash('Document deleted successfully', 'success')
    return redirect(url_for('list_documents'))

@app.route('/audit-logs')
@login_required
@admin_required
def audit_logs():
    logs = LoginLog.query.order_by(LoginLog.timestamp.desc()).all()
    return render_template('audit_logs.html', logs=logs)

def create_roles():
    admin_role = Role.query.filter_by(name='admin').first()
    user_role = Role.query.filter_by(name='user').first()
    if not admin_role:
        db.session.add(Role(name='admin'))
    if not user_role:
        db.session.add(Role(name='user'))
    db.session.commit()

    # Ensure only the first user is admin, all others are user
    users = User.query.order_by(User.id).all()
    if users:
        first_user = users[0]
        first_user.role_id = admin_role.id
        for user in users[1:]:
            user.role_id = user_role.id
        db.session.commit()

@app.route('/edit-user/<int:user_id>', methods=['GET', 'POST'])
@login_required
@admin_required
def edit_user(user_id):
    user = User.query.get_or_404(user_id)
    admin_role = Role.query.filter_by(name='admin').first()
    user_role = Role.query.filter_by(name='user').first()
    current_admin_count = User.query.join(Role).filter(Role.name == 'admin').count()
    is_only_admin = user.role_id == admin_role.id and current_admin_count == 1

    if request.method == 'POST':
        new_role = request.form.get('role')
        if new_role == 'admin':
            # Prevent more than one admin
            admin_exists = User.query.join(Role).filter(Role.name == 'admin', User.id != user.id).first() is not None
            if admin_exists:
                flash('An admin already exists. You cannot assign another admin.', 'danger')
                return redirect(url_for('edit_user', user_id=user.id))
            user.role_id = admin_role.id
        elif new_role == 'user':
            # Prevent demoting the only admin
            if is_only_admin:
                flash('You cannot demote the only admin.', 'danger')
                return redirect(url_for('edit_user', user_id=user.id))
            user.role_id = user_role.id
        db.session.commit()
        flash('User role updated successfully.', 'success')
        return redirect(url_for('manage_users'))

    return render_template('edit_user.html', user=user, is_only_admin=is_only_admin)

# Helper: Generate/load private/public key for OpenSSL digital signatures
PRIVATE_KEY_FILE = 'private_key.pem'
PUBLIC_KEY_FILE = 'public_key.pem'

def generate_rsa_key_pair():
    private_key = rsa.generate_private_key(
        public_exponent=65537,
        key_size=2048,
        backend=default_backend()
    )
    with open(PRIVATE_KEY_FILE, 'wb') as f:
        f.write(private_key.private_bytes(
            encoding=serialization.Encoding.PEM,
            format=serialization.PrivateFormat.PKCS8,
            encryption_algorithm=serialization.NoEncryption()
        ))
    with open(PUBLIC_KEY_FILE, 'wb') as f:
        f.write(private_key.public_key().public_bytes(
            encoding=serialization.Encoding.PEM,
            format=serialization.PublicFormat.SubjectPublicKeyInfo
        ))

def load_private_key():
    if not os.path.exists(PRIVATE_KEY_FILE) or not os.path.exists(PUBLIC_KEY_FILE):
        generate_rsa_key_pair()
    with open(PRIVATE_KEY_FILE, 'rb') as f:
        return serialization.load_pem_private_key(f.read(), password=None, backend=default_backend())

def load_public_key():
    with open(PUBLIC_KEY_FILE, 'rb') as f:
        return serialization.load_pem_public_key(f.read(), backend=default_backend())

def calculate_crc32(file_path):
    prev = 0
    with open(file_path, 'rb') as f:
        for line in f:
            prev = zlib.crc32(line, prev)
    return format(prev & 0xFFFFFFFF, '08x')

def sign_file_with_private_key(file_path):
    private_key = load_private_key()
    with open(file_path, 'rb') as f:
        data = f.read()
    signature = private_key.sign(
        data,
        padding.PSS(
            mgf=padding.MGF1(hashes.SHA256()),
            salt_length=padding.PSS.MAX_LENGTH
        ),
        hashes.SHA256()
    )
    return signature.hex()

def verify_file_signature(data, signature_hex):
    public_key = load_public_key()
    try:
        public_key.verify(
            bytes.fromhex(signature_hex),
            data,
            padding.PSS(
                mgf=padding.MGF1(hashes.SHA256()),
                salt_length=padding.PSS.MAX_LENGTH
            ),
            hashes.SHA256()
        )
        return True
    except Exception:
        return False

# --- Mail Setup ---
app.config['MAIL_SERVER'] = os.environ.get('MAIL_SERVER', 'smtp.gmail.com')
app.config['MAIL_PORT'] = int(os.environ.get('MAIL_PORT', 587))
app.config['MAIL_USE_TLS'] = True
app.config['MAIL_USERNAME'] = os.environ.get('MAIL_USERNAME')
app.config['MAIL_PASSWORD'] = os.environ.get('MAIL_PASSWORD')
app.config['MAIL_DEFAULT_SENDER'] = os.environ.get('MAIL_DEFAULT_SENDER', app.config['MAIL_USERNAME'])
mail = Mail(app)

# --- Password Reset Token ---
def generate_reset_token(email):
    return jwt.encode({'email': email, 'exp': datetime.utcnow() + timedelta(hours=1)}, app.secret_key, algorithm='HS256')

def verify_reset_token(token):
    try:
        data = jwt.decode(token, app.secret_key, algorithms=['HS256'])
        return data['email']
    except Exception:
        return None

@app.route('/forgot-password', methods=['GET', 'POST'])
def forgot_password():
    if request.method == 'POST':
        email = request.form['email'].strip().lower()
        user = User.query.filter_by(email=email).first()
        if user:
            token = generate_reset_token(email)
            reset_url = url_for('reset_password', token=token, _external=True)
            try:
                msg = Message('Password Reset Request', recipients=[email])
                msg.body = f"To reset your password, click the following link: {reset_url}\nIf you did not request this, ignore this email."
                mail.send(msg)
                flash('A password reset link has been sent to your email.', 'info')
            except Exception as e:
                flash('Failed to send email. Please try again later.', 'danger')
        else:
            flash('Email not found.', 'danger')
        return redirect(url_for('login'))
    return render_template('forgot_password.html')

@app.route('/reset-password/<token>', methods=['GET', 'POST'])
def reset_password(token):
    email = verify_reset_token(token)
    if not email:
        flash('Invalid or expired reset link.', 'danger')
        return redirect(url_for('login'))
    if request.method == 'POST':
        password = request.form['password']
        if not valid_password(password):
            flash('Password too weak. Must be at least 8 characters, include numbers and symbols.', 'danger')
            return render_template('reset_password.html', token=token)
        user = User.query.filter_by(email=email).first()
        if user:
            user.password_hash = bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
            db.session.commit()
            flash('Password reset successful. Please log in.', 'success')
            return redirect(url_for('login'))
        else:
            flash('User not found.', 'danger')
            return redirect(url_for('login'))
    return render_template('reset_password.html', token=token)

@app.route('/change-password', methods=['GET', 'POST'])
@login_required
def change_password():
    if request.method == 'POST':
        old_password = request.form['old_password']
        new_password = request.form['new_password']
        if not bcrypt.checkpw(old_password.encode('utf-8'), current_user.password_hash.encode('utf-8')):
            flash('Old password is incorrect.', 'danger')
            return render_template('change_password.html')
        if not valid_password(new_password):
            flash('Password too weak. Must be at least 8 characters, include numbers and symbols.', 'danger')
            return render_template('change_password.html')
        current_user.password_hash = bcrypt.hashpw(new_password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
        db.session.commit()
        flash('Password changed successfully.', 'success')
        return redirect(url_for('profile'))
    return render_template('change_password.html')

@app.route('/profile', methods=['GET', 'POST'])
@login_required
def profile():
    if request.method == 'POST':
        username = request.form.get('username', '').strip()
        phone = request.form.get('phone', '').strip()
        department = request.form.get('department', '').strip()
        job_title = request.form.get('job_title', '').strip()
        profile_pic = request.files.get('profile_pic')
        # Validate and update fields
        if username:
            current_user.username = username
        current_user.phone = phone
        current_user.department = department
        current_user.job_title = job_title
        # Handle profile picture upload
        if profile_pic and profile_pic.filename:
            allowed_ext = {'jpg', 'jpeg', 'png', 'gif'}
            ext = profile_pic.filename.rsplit('.', 1)[-1].lower()
            if ext not in allowed_ext:
                flash('Image type not allowed. Only JPG, PNG, GIF.', 'danger')
                return render_template('profile.html', user=current_user)
            if len(profile_pic.read()) > 5 * 1024 * 1024:
                flash('Image too large. Max size is 5MB.', 'danger')
                return render_template('profile.html', user=current_user)
            profile_pic.seek(0)
            filename = f"user_{current_user.id}_profile.{ext}"
            save_path = os.path.join(app.static_folder, 'profile_pics')
            os.makedirs(save_path, exist_ok=True)
            file_path = os.path.join(save_path, filename)
            profile_pic.save(file_path)
            current_user.profile_pic_url = url_for('static', filename=f'profile_pics/{filename}')
        db.session.commit()
        flash('Profile updated successfully.', 'success')
    return render_template('profile.html', user=current_user)

@app.route('/edit-document/<int:doc_id>', methods=['GET', 'POST'])
@login_required
@admin_required
def edit_document(doc_id):
    document = Document.query.get_or_404(doc_id)
    
    if request.method == 'POST':
        description = request.form.get('description', '').strip()
        document.description = description
        document.last_modified = datetime.utcnow()
        db.session.commit()
        flash('Document description updated successfully.', 'success')
        return redirect(url_for('list_documents'))
    
    return render_template('edit_document.html', document=document)

@app.route('/add-user', methods=['GET', 'POST'])
@login_required
@admin_required
def add_user():
    if request.method == 'POST':
        username = request.form['username'].strip()
        email = request.form['email'].strip().lower()
        password = request.form['password']
        is_admin = 'is_admin' in request.form

        if not username or not email or not password:
            flash('All fields are required.', 'danger')
            return redirect(url_for('add_user'))
        
        if not valid_password(password):
            flash('Password too weak. Must be at least 8 characters, include numbers and symbols.', 'danger')
            return redirect(url_for('add_user'))
        
        if User.query.filter_by(email=email).first():
            flash('Email already used, try another.', 'danger')
            return redirect(url_for('add_user'))
        
        hashed = bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt())
        
        admin_role = Role.query.filter_by(name='admin').first()
        user_role = Role.query.filter_by(name='user').first()
        
        # Check if we're trying to create an admin
        if is_admin:
            admin_exists = User.query.join(Role).filter(Role.name == 'admin').first() is not None
            if admin_exists:
                flash('An admin already exists. You cannot create another admin.', 'danger')
                return redirect(url_for('add_user'))
            role_id = admin_role.id
        else:
            role_id = user_role.id

        user = User(
            username=username,
            email=email,
            password_hash=hashed.decode('utf-8'),
            auth_method='manual',
            role_id=role_id
        )

        try:
            db.session.add(user)
            db.session.commit()
            save_user_to_file(user)
            flash('User added successfully.', 'success')
            return redirect(url_for('manage_users'))
        except IntegrityError:
            db.session.rollback()
            flash('Username or email already exists.', 'danger')
    
    return render_template('add_user.html')

@app.route('/view-document/<int:doc_id>')
@login_required
@require_2fa
def view_document(doc_id):
    document = Document.query.get_or_404(doc_id)
    # Check permissions
    if not current_user.is_admin() and document.user_id != current_user.id:
        abort(403)
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], document.filename)
    # Decrypt the file
    decrypted_data = decrypt_file(file_path)
    ext = document.original_filename.rsplit('.', 1)[-1].lower()
    if ext == 'pdf':
        return send_file(
            io.BytesIO(decrypted_data),
            mimetype='application/pdf',
            as_attachment=False,
            download_name=document.original_filename
        )
    elif ext == 'txt':
        return send_file(
            io.BytesIO(decrypted_data),
            mimetype='text/plain',
            as_attachment=False,
            download_name=document.original_filename
        )
    elif ext == 'docx':
        # Try to extract text from docx
        try:
            import docx
            from tempfile import NamedTemporaryFile
            with NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                tmp.write(decrypted_data)
                tmp.flush()
                doc = docx.Document(tmp.name)
                text = '\n'.join([p.text for p in doc.paragraphs])
            os.unlink(tmp.name)
            return f'<pre>{text}</pre>'
        except Exception as e:
            return f'Cannot preview DOCX file: {e}'
    else:
        return 'Preview not supported for this file type.'

def get_simple_file_type(filename):
    ext = filename.rsplit('.', 1)[-1].lower()
    if ext in ['pdf', 'txt', 'docx']:
        return ext
    return 'other'

if __name__ == '__main__':
    with app.app_context():
        create_roles()
    app.run(host='127.0.0.1', port=5000, debug=True, ssl_context='adhoc')